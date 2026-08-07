#!/usr/bin/env python3
"""
CDE Converter microservice — port 5001
Cross-platform: Windows + Ubuntu/Linux/Mac
"""

import os, sys, io, json, shutil, tempfile, subprocess, traceback, time, platform
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse
from pathlib import Path

IS_WINDOWS = platform.system() == "Windows"

# ── ezdxf ──────────────────────────────────────────────────────
try:
    import ezdxf
    from ezdxf.addons.drawing import RenderContext, Frontend
    from ezdxf.addons.drawing.svg import SVGBackend
    from ezdxf.addons.drawing.properties import LayoutProperties
    EZDXF_OK = True
except ImportError as e:
    EZDXF_OK = False
    EZDXF_ERR = str(e)

PORT = int(os.environ.get("CONVERTER_PORT", 5001))

DWG_VERSIONS = {
    "AC1015":"AutoCAD 2000","AC1018":"AutoCAD 2004","AC1021":"AutoCAD 2007",
    "AC1024":"AutoCAD 2010","AC1027":"AutoCAD 2013","AC1032":"AutoCAD 2018",
    "AC1035":"AutoCAD 2021","AC1037":"AutoCAD 2023",
}

OFFICE_EXTS = {
    "doc","docx","xls","xlsx","ppt","pptx",
    "odt","ods","odp","rtf","txt","csv","html","htm"
}

# ══════════════════════════════════════════════════════════════
#  TEMP DIR — use short path on Windows to avoid spaces
# ══════════════════════════════════════════════════════════════
def make_temp_dir():
    """
    On Windows, tempfile defaults to C:\\Users\\...\\AppData\\Local\\Temp
    which has spaces. Use C:\\Temp instead if it exists or can be created,
    as short paths are safer for legacy tools like ODA.
    """
    if IS_WINDOWS:
        base = "C:\\Temp"
        try:
            os.makedirs(base, exist_ok=True)
            return tempfile.mkdtemp(dir=base)
        except Exception:
            pass  # fall through to default
    return tempfile.mkdtemp()


def safe_rmtree(path):
    try:
        shutil.rmtree(path, ignore_errors=True)
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════
#  SUBPROCESS — Windows-safe, handles spaces in paths
# ══════════════════════════════════════════════════════════════
def run_cmd(cmd, timeout=120):
    """
    Run command list. On Windows uses shell=False with proper quoting.
    All paths in cmd should already be absolute strings.
    Returns (returncode, stdout, stderr).
    """
    log_cmd = " ".join(f'"{c}"' if " " in str(c) else str(c) for c in cmd)
    print(f"[CMD] {log_cmd}", flush=True)

    kwargs = dict(timeout=timeout)

    if IS_WINDOWS:
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        si.wShowWindow = 0  # SW_HIDE
        kwargs["startupinfo"] = si
        kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
        # On Windows, pass cmd as list (Python handles quoting internally)
        kwargs["shell"] = False

    try:
        r = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            **kwargs
        )
        stdout = r.stdout.decode("utf-8", errors="replace") if r.stdout else ""
        stderr = r.stderr.decode("utf-8", errors="replace") if r.stderr else ""
        print(f"[CMD] rc={r.returncode} stdout={stdout[:150]} stderr={stderr[:150]}", flush=True)
        return r.returncode, stdout, stderr
    except subprocess.TimeoutExpired:
        print("[CMD] TIMEOUT", flush=True)
        return -1, "", f"Timed out after {timeout}s"
    except FileNotFoundError as e:
        print(f"[CMD] NOT FOUND: {e}", flush=True)
        return -2, "", f"Executable not found: {e}"
    except Exception as e:
        print(f"[CMD] ERROR: {e}", flush=True)
        return -3, "", str(e)


# ══════════════════════════════════════════════════════════════
#  LIBREOFFICE
# ══════════════════════════════════════════════════════════════
def find_libreoffice():
    if IS_WINDOWS:
        for base in [r"C:\Program Files", r"C:\Program Files (x86)"]:
            if not os.path.isdir(base):
                continue
            for entry in sorted(os.listdir(base), reverse=True):
                if "libreoffice" in entry.lower():
                    exe = os.path.join(base, entry, "program", "soffice.exe")
                    if os.path.isfile(exe):
                        return exe
        # Explicit fallbacks
        for p in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            if os.path.isfile(p): return p
    else:
        for cmd in ["libreoffice", "soffice"]:
            found = shutil.which(cmd)
            if found: return found
        mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.isfile(mac): return mac
    return None


def convert_office_to_pdf(file_path: str) -> dict:
    lo = find_libreoffice()
    if not lo:
        return {"success": False, "error":
            "LibreOffice not installed.\n"
            "Ubuntu: sudo apt install libreoffice\n"
            "Windows: https://www.libreoffice.org/download"}

    abs_path = str(Path(file_path).resolve())
    out_dir = make_temp_dir()
    try:
        cmd = [lo, "--headless", "--norestore", "--nofirststartwizard",
               "--convert-to", "pdf", "--outdir", out_dir, abs_path]
        rc, stdout, stderr = run_cmd(cmd, timeout=90)

        # LibreOffice on Windows prints harmless warnings to stderr — ignore them
        # Check for actual output files
        all_files = list(Path(out_dir).iterdir())
        print(f"[LO] out_dir={out_dir} files={all_files}", flush=True)

        # Use rglob to find PDF anywhere in out_dir (LO sometimes makes subfolders)
        pdf_files = list(set(Path(out_dir).rglob("*.pdf")))
        if not pdf_files:
            # LO may have written to the SOURCE file's directory instead of out_dir
            # This happens when --outdir is ignored (rare LO bug on Windows)
            src_dir = Path(abs_path).parent
            fallback = list(src_dir.glob(Path(abs_path).stem + "*.pdf"))
            print(f"[LO] Fallback PDF search in {src_dir}: {fallback}", flush=True)
            if fallback:
                return {"success": True, "pdfBytes": fallback[0].read_bytes()}
            return {"success": False, "error":
                f"LibreOffice produced no PDF (rc={rc}).\n"
                f"stdout={stdout[:300]}\nstderr={stderr[:200]}"}

        print(f"[LO] PDF found: {pdf_files[0]}", flush=True)
        return {"success": True, "pdfBytes": pdf_files[0].read_bytes()}
    finally:
        safe_rmtree(out_dir)


# ══════════════════════════════════════════════════════════════
#  ODA FILE CONVERTER
# ══════════════════════════════════════════════════════════════
def find_oda():
    if IS_WINDOWS:
        for base in [r"C:\Program Files\ODA", r"C:\Program Files (x86)\ODA"]:
            if not os.path.isdir(base):
                continue
            for entry in sorted(os.listdir(base), reverse=True):
                exe = os.path.join(base, entry, "ODAFileConverter.exe")
                if os.path.isfile(exe):
                    return exe
    for p in ["/usr/bin/ODAFileConverter", "/usr/local/bin/ODAFileConverter",
              os.environ.get("ODA_PATH", "")]:
        if p and os.path.isfile(p):
            return p
    return shutil.which("ODAFileConverter")


def dwg_via_oda(dwg_path: str) -> dict:
    oda = find_oda()
    if not oda:
        return {"success": False, "error": "ODA_NOT_FOUND"}

    abs_dwg = str(Path(dwg_path).resolve())

    # Use short C:\Temp paths to avoid spaces breaking ODA
    in_dir  = make_temp_dir()
    out_dir = make_temp_dir()

    try:
        dest = os.path.join(in_dir, Path(abs_dwg).name)
        shutil.copy2(abs_dwg, dest)
        print(f"[ODA] Copied DWG to: {dest}", flush=True)
        print(f"[ODA] Output dir:    {out_dir}", flush=True)
        print(f"[ODA] Executable:    {oda}", flush=True)

        logs = []
        for version in ["ACAD2018", "ACAD2013", "ACAD2010", "ACAD2007"]:
            # ODA CLI: <inputDir> <outputDir> <version> <type> <recurse> <audit>
            # Pass as list — Python/Windows handles quoting for spaces automatically
            cmd = [oda, in_dir, out_dir, version, "DXF", "0", "1"]
            rc, stdout, stderr = run_cmd(cmd, timeout=120)
            logs.append(f"v={version} rc={rc} err={stderr[:150]}")

            # ODA can still be writing — wait for flush
            time.sleep(2)

            # Deduplicate (Windows rglob returns same file twice for *.dxf + *.DXF)
            dxf_files = list({str(f): f for f in
                (list(Path(out_dir).rglob("*.dxf")) +
                 list(Path(out_dir).rglob("*.DXF")))}.values())
            print(f"[ODA] DXF files after {version}: {dxf_files}", flush=True)

            if dxf_files:
                dxf_path = dxf_files[0]
                print(f"[ODA] Reading DXF: {dxf_path} ({dxf_path.stat().st_size} bytes)", flush=True)
                content = dxf_path.read_text(encoding="utf-8", errors="replace")
                print(f"[ODA] DXF content preview: {content[:120]!r}", flush=True)
                result = render_dxf_string(content)
                if not result.get("success"):
                    print(f"[ODA] ezdxf render failed: {result.get('error','?')}", flush=True)
                result["odaLog"] = "\n".join(logs)
                result["odaVersion"] = version
                return result

        return {"success": False,
                "error": "ODA ran but produced no DXF output.\n" + "\n".join(logs)}
    finally:
        safe_rmtree(in_dir)
        safe_rmtree(out_dir)


# ══════════════════════════════════════════════════════════════
#  LibreDWG
# ══════════════════════════════════════════════════════════════
def find_dwg2dxf():
    for p in ["/usr/bin/dwg2dxf", "/usr/local/bin/dwg2dxf",
              r"C:\Program Files\LibreDWG\dwg2dxf.exe",
              os.environ.get("DWG2DXF_PATH", "")]:
        if p and os.path.isfile(p): return p
    return shutil.which("dwg2dxf")


def find_tesseract():
    """
    Locate the Tesseract OCR binary. Single source of truth — the same
    discovery was previously inlined in three places (text extraction,
    startup banner, health check) with slightly different candidate lists.
    """
    found = shutil.which("tesseract")
    if found:
        return found
    if IS_WINDOWS:
        for p in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(
                os.environ.get("USERNAME", "")),
        ]:
            if os.path.isfile(p):
                return p
    p = os.environ.get("TESSERACT_PATH", "")
    return p if p and os.path.isfile(p) else None


def dwg_via_libredwg(dwg_path: str) -> dict:
    tool = find_dwg2dxf()
    if not tool:
        return {"success": False, "error": "LIBREDWG_NOT_FOUND"}
    abs_dwg = str(Path(dwg_path).resolve())
    tmp = make_temp_dir()
    try:
        out_dxf = os.path.join(tmp, "out.dxf")
        rc, stdout, stderr = run_cmd([tool, "-o", out_dxf, abs_dwg], timeout=60)
        if not os.path.exists(out_dxf):
            return {"success": False,
                    "error": f"dwg2dxf no output (rc={rc}). {stderr[:200]}"}
        return render_dxf_string(Path(out_dxf).read_text(encoding="utf-8", errors="replace"))
    finally:
        safe_rmtree(tmp)


# ══════════════════════════════════════════════════════════════
#  ezdxf DXF -> SVG
# ══════════════════════════════════════════════════════════════
def render_dxf(path: str) -> dict:
    try:
        return render_dxf_string(Path(path).read_text(encoding="utf-8", errors="replace"))
    except Exception as e:
        return {"success": False, "error": f"Cannot read DXF: {e}"}


def render_dxf_string(dxf_content: str) -> dict:
    if not EZDXF_OK:
        return {"success": False, "error": f"ezdxf not installed: {EZDXF_ERR}"}
    # Write to temp file — avoids variable-shadowing with ezdxf module name
    tmp_dxf = None
    try:
        with tempfile.NamedTemporaryFile(mode="w", suffix=".dxf",
                                         encoding="utf-8", delete=False) as f:
            f.write(dxf_content)
            tmp_dxf = f.name
        return _render_dxf_file(tmp_dxf)
    finally:
        if tmp_dxf and os.path.exists(tmp_dxf):
            try: os.unlink(tmp_dxf)
            except: pass


def _render_dxf_file(dxf_path: str) -> dict:
    """Render a DXF file to SVG using ezdxf with recover fallback."""
    import ezdxf as _ezdxf
    import ezdxf.recover as _recover

    doc = None
    try:
        doc = _ezdxf.readfile(dxf_path)
        print(f"[ezdxf] Normal read OK", flush=True)
    except Exception as e1:
        print(f"[ezdxf] Normal read failed: {e1} — trying recover", flush=True)
        try:
            doc, auditor = _recover.read(dxf_path)
            if auditor.has_errors:
                print(f"[ezdxf] Recover fixed {len(auditor.errors)} errors", flush=True)
            print(f"[ezdxf] Recover read OK", flush=True)
        except Exception as e2:
            print(f"[ezdxf] Recover also failed: {e2}", flush=True)
            return {"success": False, "error": f"ezdxf cannot parse DXF: {e2}"}

    try:
        # ezdxf 1.4.4 API
        from ezdxf.addons.drawing import RenderContext, Frontend
        from ezdxf.addons.drawing.svg import SVGBackend
        from ezdxf.addons.drawing.layout import Page, Settings
        from ezdxf.addons.drawing.properties import LayoutProperties

        msp = doc.modelspace()
        entity_count = len(list(msp))
        print(f"[ezdxf] modelspace entities: {entity_count}", flush=True)

        # If modelspace empty, try paper space layouts
        layout_to_render = msp
        if entity_count == 0:
            layouts = [l for l in doc.layouts if l.name != "Model"]
            print(f"[ezdxf] Modelspace empty — checking {len(layouts)} paper layouts", flush=True)
            for layout in layouts:
                lcount = len(list(layout))
                print(f"[ezdxf]   layout '{layout.name}': {lcount} entities", flush=True)
                if lcount > 0:
                    layout_to_render = layout
                    entity_count = lcount
                    break

        # ezdxf 1.x: SVGBackend() takes no args; output via get_string(page)
        backend = SVGBackend()
        lp = LayoutProperties.from_layout(layout_to_render)
        lp.set_colors(bg="#1a1d27")
        Frontend(RenderContext(doc), backend).draw_layout(
            layout_to_render, finalize=True, layout_properties=lp)

        page = Page(420, 297)                   # A4 landscape mm
        settings = Settings(fit_page=True)
        svg = backend.get_string(page, settings=settings)

        print(f"[ezdxf] SVG length: {len(svg)}", flush=True)
        if not svg or "<svg" not in svg:
            return {"success": False,
                    "error": f"ezdxf rendered empty SVG. Entities: {entity_count}, "
                             f"version: {doc.dxfversion}. Layers may be frozen/off."}

        # Convert mm dimensions to pixels (96dpi) so browser renders at correct size
        # ezdxf outputs: width="420mm" height="297mm" — replace with px values
        import re
        def mm_to_px(mm): return round(float(mm) * 3.7795275591)
        def replace_dim(m):
            val = m.group(1); unit = m.group(2)
            if unit == 'mm': return f'{mm_to_px(val)}px'
            if unit == 'cm': return f'{mm_to_px(float(val)*10)}px'
            return m.group(0)
        svg = re.sub(r'([\d.]+)(mm|cm)', replace_dim, svg, count=4)
        svg = svg.replace("<svg ",
            '<svg style="max-width:100%;max-height:100%;display:block" ', 1)
        return {"success": True, "svg": svg,
                "entityCount": entity_count, "dxfVersion": doc.dxfversion}
    except Exception as e:
        print(f"[ezdxf] render error:\n{traceback.format_exc()}", flush=True)
        return {"success": False, "error": f"ezdxf render error: {e}"}



# ══════════════════════════════════════════════════════════════
#  DWG magic byte detection
# ══════════════════════════════════════════════════════════════
def is_dwg(path: str) -> tuple:
    try:
        with open(path, "rb") as f:
            magic = f.read(6).decode("ascii", errors="replace")
        if magic.startswith("AC"):
            return True, DWG_VERSIONS.get(magic[:6], magic[:6])
    except Exception:
        pass
    return False, ""


# ══════════════════════════════════════════════════════════════
#  Main dispatcher
# ══════════════════════════════════════════════════════════════
def is_revit(path: str) -> bool:
    """Detect Revit .rvt/.rfa files by OLE2 magic bytes."""
    try:
        with open(path, "rb") as f:
            magic = f.read(8)
        return magic[:4] == bytes([0xD0, 0xCF, 0x11, 0xE0])
    except Exception:
        return False


def ifc_to_gltf_json(ifc_path: str) -> dict:
    """
    Convert IFC file to a glTF-compatible JSON structure containing
    mesh data (vertices, faces, colors per element type) for Three.js.
    Returns {"success": True, "gltfData": {...}} or {"success": False, "error": ...}
    """
    try:
        import ifcopenshell
        import ifcopenshell.geom
        import numpy as np
        import json, base64, struct
    except ImportError as e:
        return {"success": False, "error": f"ifcopenshell not installed: {e}\nRun: pip install ifcopenshell"}

    try:
        ifc = ifcopenshell.open(ifc_path)
        print(f"[IFC] Opened: schema={ifc.schema}, products={len(ifc.by_type('IfcProduct'))}", flush=True)
    except Exception as e:
        return {"success": False, "error": f"Cannot open IFC: {e}"}

    # Colour map by IFC type
    TYPE_COLORS = {
        "IfcWall":          [0.85, 0.82, 0.78, 1.0],
        "IfcWallStandardCase": [0.85, 0.82, 0.78, 1.0],
        "IfcSlab":          [0.75, 0.75, 0.75, 1.0],
        "IfcRoof":          [0.62, 0.45, 0.35, 1.0],
        "IfcColumn":        [0.80, 0.75, 0.70, 1.0],
        "IfcBeam":          [0.70, 0.65, 0.60, 1.0],
        "IfcDoor":          [0.65, 0.45, 0.25, 1.0],
        "IfcWindow":        [0.55, 0.75, 0.90, 0.5],
        "IfcStair":         [0.80, 0.78, 0.75, 1.0],
        "IfcRamp":          [0.78, 0.76, 0.72, 1.0],
        "IfcFurnishingElement": [0.60, 0.50, 0.40, 1.0],
        "IfcFlowTerminal":  [0.40, 0.65, 0.80, 1.0],
        "IfcFlowSegment":   [0.70, 0.70, 0.30, 1.0],
        "IfcPlate":         [0.75, 0.73, 0.70, 1.0],
        "IfcMember":        [0.65, 0.60, 0.55, 1.0],
    }
    DEFAULT_COLOR = [0.70, 0.68, 0.65, 1.0]

    settings = ifcopenshell.geom.settings()
    settings.set(settings.USE_WORLD_COORDS, True)
    settings.set(settings.WELD_VERTICES, True)

    all_positions = []
    all_normals   = []
    all_indices   = []
    all_colors    = []
    vertex_offset = 0
    mesh_count    = 0

    try:
        # Filter to only element types that have geometry
        include_types = [
            'IfcWall','IfcWallStandardCase','IfcSlab','IfcRoof','IfcColumn',
            'IfcBeam','IfcDoor','IfcWindow','IfcStair','IfcRamp',
            'IfcFurnishingElement','IfcPlate','IfcMember','IfcCovering',
            'IfcFlowTerminal','IfcFlowSegment','IfcOpeningElement',
        ]
        # Try with type filter first, fall back to all types
        it = None
        for attempt in ['filtered', 'all']:
            try:
                if attempt == 'filtered':
                    it = ifcopenshell.geom.iterator(settings, ifc, include_entities=include_types)
                else:
                    it = ifcopenshell.geom.iterator(settings, ifc)
                if it.initialize():
                    print(f"[IFC] Iterator initialized ({attempt})", flush=True)
                    break
                else:
                    it = None
            except Exception as e:
                print(f"[IFC] Iterator ({attempt}) error: {e}", flush=True)
                it = None

        if it is None:
            prod_count = len(ifc.by_type('IfcProduct'))
            return {"success": False, "error":
                f"IFC file has no renderable geometry ({prod_count} products found). "
                f"The file may contain only 2D data or have no geometric representations."}

        while True:
            shape = it.get()
            geo   = shape.geometry
            verts  = np.array(geo.verts,  dtype=np.float32).reshape(-1, 3)
            norms  = np.array(geo.normals,dtype=np.float32).reshape(-1, 3) if geo.normals else np.zeros_like(verts)
            faces  = np.array(geo.faces,  dtype=np.uint32).reshape(-1, 3)

            if len(verts) == 0 or len(faces) == 0:
                if not it.next(): break
                continue

            col = TYPE_COLORS.get(shape.type, DEFAULT_COLOR)
            colors = np.tile(col[:3], (len(verts), 1)).astype(np.float32)

            all_positions.append(verts)
            all_normals.append(norms)
            all_indices.append(faces + vertex_offset)
            all_colors.append(colors)
            vertex_offset += len(verts)
            mesh_count += 1

            if not it.next():
                break

    except Exception as e:
        return {"success": False, "error": f"Geometry extraction failed: {e}"}

    if mesh_count == 0:
        return {"success": False, "error": "No geometry found in IFC file. "
                "Ensure the file contains IfcWall, IfcSlab or other building elements with geometry."}

    # Concatenate all geometry
    positions = np.concatenate(all_positions, axis=0).astype(np.float32)
    normals   = np.concatenate(all_normals,   axis=0).astype(np.float32)
    indices   = np.concatenate(all_indices,   axis=0).astype(np.uint32)
    colors    = np.concatenate(all_colors,    axis=0).astype(np.float32)

    print(f"[IFC] Total: {len(positions)} verts, {len(indices)} triangles, {mesh_count} elements", flush=True)

    def to_b64(arr):
        return base64.b64encode(arr.tobytes()).decode()

    # Build minimal glTF-compatible data packet for Three.js
    gltf_data = {
        "positions": to_b64(positions),
        "normals":   to_b64(normals),
        "indices":   to_b64(indices),
        "colors":    to_b64(colors),
        "vertexCount":   int(len(positions)),
        "triangleCount": int(len(indices)),
        "elementCount":  int(mesh_count),
        "schema":        ifc.schema,
        "bounds": {
            "min": positions.min(axis=0).tolist(),
            "max": positions.max(axis=0).tolist(),
        }
    }

    return {"success": True, "type": "ifc3d", "gltfData": gltf_data}


def convert(file_path: str, content_type: str = "") -> dict:
    # Resolve to absolute, normalise slashes
    try:
        abs_path = str(Path(file_path).resolve())
    except Exception as e:
        return {"success": False, "error": f"Invalid path: {file_path} — {e}"}

    print(f"[CONVERT] path={abs_path} ct={content_type}", flush=True)

    if not os.path.exists(abs_path):
        return {"success": False,
                "error": f"File not found: {abs_path}\n"
                         f"(received from Java: {file_path})"}

    ext = Path(abs_path).suffix.lower().lstrip(".")
    ct  = content_type.lower()

    # Office -> PDF
    if ext in OFFICE_EXTS or any(x in ct for x in [
            "word","excel","powerpoint","opendocument","rtf","text/plain","text/csv"]):
        r = convert_office_to_pdf(abs_path)
        if r["success"]:
            return {"success": True, "type": "pdf", "pdfBytes": r["pdfBytes"]}
        return {"success": False, "error": r["error"], "type": "office_error",
                "loInstalled": find_libreoffice() is not None}

    # PDF passthrough
    if ext == "pdf" or "pdf" in ct:
        return {"success": True, "type": "pdf",
                "pdfBytes": Path(abs_path).read_bytes()}

    # DWG -> DXF -> SVG
    dwg, version = is_dwg(abs_path)
    if dwg or ext == "dwg":
        r = dwg_via_oda(abs_path)
        if r.get("success"):
            r["convertedBy"] = "ODA"; r["dwgVersion"] = version; return r
        oda_err = r.get("error", "")

        r = dwg_via_libredwg(abs_path)
        if r.get("success"):
            r["convertedBy"] = "LibreDWG"; r["dwgVersion"] = version; return r

        return {"success": False, "error": "DWG_NEED_CONVERTER",
                "dwgVersion": version,
                "odaInstalled": find_oda() is not None,
                "libredwgInstalled": find_dwg2dxf() is not None,
                "odaError": oda_err,
                "libredwgError": r.get("error", "")}

    # DXF -> SVG
    if ext == "dxf" or "dxf" in ct:
        return render_dxf(abs_path)

    # 3D formats — also detect by content type
    if ext == "ifc" or "ifc" in ct or "step" in ct:
        print(f"[IFC] Routing to ifc_to_gltf_json: {abs_path}", flush=True)
        return ifc_to_gltf_json(abs_path)

    if ext in ("rvt", "rfa"):
        return {
            "success": False, "error": "REVIT_BINARY",
            "type": "revit_binary",
            "fileName": os.path.basename(abs_path)
        }

    if ext in ("glb", "gltf", "obj", "stl", "ply", "dae", "3ds"):
        return {
            "success": True, "type": "model3d_passthrough",
            "ext": ext,
            "filePath": abs_path
        }

    return {"success": False, "error": f"Unsupported type: .{ext}"}


# ══════════════════════════════════════════════════════════════
#  HTTP Server
# ══════════════════════════════════════════════════════════════
class Handler(BaseHTTPRequestHandler):

    def log_message(self, fmt, *args):
        if args and str(args[1]) not in ("200", "204"):
            super().log_message(fmt, *args)

    def do_GET(self):
        if urlparse(self.path).path == "/health":
            lo   = find_libreoffice()
            tess = find_tesseract()
            self._json(200, {
                "status": "ok", "platform": platform.system(),
                "ezdxf": EZDXF_OK,
                "ezdxfVersion": ezdxf.__version__ if EZDXF_OK else None,
                "libreoffice": lo is not None, "libreofficePath": lo,
                "odaInstalled": find_oda() is not None, "odaPath": find_oda(),
                "libredwgInstalled": find_dwg2dxf() is not None,
                "tesseractInstalled": tess is not None, "tesseractPath": tess,
            })
        else:
            self._json(404, {"error": "Use POST /convert"})

    def do_POST(self):
        ppath  = urlparse(self.path).path
        length = int(self.headers.get("Content-Length", 0))
        raw    = self.rfile.read(length) if length else b""
        try:
            body = json.loads(raw) if raw else {}
        except Exception:
            self._json(400, {"success": False, "error": "Invalid JSON"}); return

        if ppath == "/redact":
            result = redact_pdf(
                body.get("path",""),
                body.get("regions", []),
                body.get("output",""),
                body.get("burn", True)
            )
            self._json(200, result)
            return

        if ppath == "/ocr":
            result = ocr_pdf_to_searchable(
                body.get("path",""),
                body.get("output",""),
                body.get("lang","eng"),
                body.get("dpi",300),
                body.get("skipTextPages", True)
            )
            self._json(200, result)
            return

        if ppath == "/flatten":
            result = flatten_annotations_to_pdf(
                body.get("path",""),
                body.get("shapes", []),
                body.get("output",""),
                body.get("quality","screen")
            )
            self._json(200, result)
            return

        if ppath == "/form-fields":
            result = inspect_pdf_form(body.get("path",""))
            self._json(200, result)
            return

        if ppath == "/form-fill":
            result = fill_pdf_form(
                body.get("path",""),
                body.get("fields", {}),
                body.get("output",""),
                body.get("flatten", False)
            )
            self._json(200, result)
            return

        if ppath == "/ifc-tree":
            result = extract_ifc_tree(body.get("path",""))
            self._json(200, result)
            return

        if ppath == "/compare":
            path1 = body.get("path1","").strip()
            path2 = body.get("path2","").strip()
            ct1   = body.get("contentType1","")
            ct2   = body.get("contentType2","")
            if not path1 or not path2:
                self._json(400, {"success": False, "error": "Missing path1 or path2"}); return
            self._json(200, compare_files(path1, path2, ct1, ct2))
            return

        if ppath != "/convert":
            self._json(404, {"error": "Not found"}); return

        file_path    = body.get("path", "").strip()
        content_type = body.get("contentType", "")

        if not file_path:
            self._json(400, {"success": False, "error": "Missing 'path'"}); return

        result = convert(file_path, content_type)

        rtype = result.get("type", "")

        # PDF — send as binary
        if result.get("success") and rtype == "pdf":
            pdf_bytes = result["pdfBytes"]
            self.send_response(200)
            self.send_header("Content-Type", "application/pdf")
            self.send_header("Content-Length", str(len(pdf_bytes)))
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Disposition", "inline")
            self.end_headers()
            self.wfile.write(pdf_bytes)

        # 3D model passthrough (GLB/OBJ/STL) — serve raw bytes
        elif result.get("success") and rtype == "model3d_passthrough":
            ext = result.get("ext","glb")
            mime_map = {"glb":"model/gltf-binary","gltf":"model/gltf+json",
                        "obj":"text/plain","stl":"application/octet-stream",
                        "ply":"application/octet-stream","dae":"text/xml"}
            mime = mime_map.get(ext, "application/octet-stream")
            raw = Path(result["filePath"]).read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Length", str(len(raw)))
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(raw)

        # IFC 3D data — JSON with base64 geometry
        elif result.get("success") and rtype == "ifc3d":
            result.pop("pdfBytes", None)
            self._json(200, result)

        else:
            result.pop("pdfBytes", None)
            self._json(200, result)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, GET, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def _json(self, status, data):
        body = json.dumps(data).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(body)



# ══════════════════════════════════════════════════════════════
#  3D VIEWER — IFC / glTF / OBJ / STL
# ══════════════════════════════════════════════════════════════

MODEL_3D_EXTS = {"ifc", "glb", "gltf", "obj", "stl", "ply", "dae", "3ds"}



# ══════════════════════════════════════════════════════════════
#  FILE COMPARISON ENGINE
# ══════════════════════════════════════════════════════════════

def compare_files(path1: str, path2: str, ct1: str = "", ct2: str = "") -> dict:
    """
    Compare two files and return structured change report.
    Supports: DXF/DWG, IFC, PDF/Office (text), images
    """
    p1 = Path(path1); p2 = Path(path2)
    ext1 = p1.suffix.lower().lstrip('.')
    ext2 = p2.suffix.lower().lstrip('.')

    print(f"[COMPARE] {p1.name} vs {p2.name}", flush=True)

    # Route by file type
    if ext1 in ('dxf', 'dwg') and ext2 in ('dxf', 'dwg'):
        return _compare_cad(path1, path2, ext1, ext2)
    elif ext1 == 'ifc' and ext2 == 'ifc':
        return _compare_ifc(path1, path2)
    elif ext1 in OFFICE_EXTS | {'pdf'} and ext2 in OFFICE_EXTS | {'pdf'}:
        return _compare_documents(path1, path2, ext1, ext2)
    elif ext1 in ('png','jpg','jpeg','bmp','gif') and ext2 in ('png','jpg','jpeg','bmp','gif'):
        return _compare_images(path1, path2)
    else:
        # Mixed types — do basic metadata comparison
        return _compare_metadata(path1, path2)


# ── CAD Comparison (DXF/DWG) ──────────────────────────────────
def _load_dxf_for_compare(path: str, ext: str):
    """Load a DXF or DWG file, converting DWG via ODA if needed."""
    if ext == 'dwg':
        # Convert to DXF first
        r = dwg_via_oda(path)
        if not r.get('success'):
            r2 = dwg_via_libredwg(path)
            if not r2.get('success'):
                return None, f"Cannot read DWG: {r.get('error','')}"
        # Write DXF to temp file
        import ezdxf as _ezdxf
        tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.dxf',
                                          encoding='utf-8', delete=False)
        # DWG->DXF was done inside the ODA func, need to re-read the DXF content
        # For simplicity, re-run ODA and get the DXF path directly
        return _load_dxf_from_path_via_oda(path)
    else:
        try:
            import ezdxf as _ezdxf
            doc = _ezdxf.readfile(path)
            return doc, None
        except Exception as e:
            try:
                import ezdxf.recover as _recover
                doc, _ = _recover.read(path)
                return doc, None
            except Exception as e2:
                return None, str(e2)

def _load_dxf_from_path_via_oda(dwg_path: str):
    """Returns (ezdxf_doc, error_str)"""
    import ezdxf as _ezdxf, ezdxf.recover as _recover
    oda = find_oda()
    if not oda:
        return None, "ODA not installed"
    in_dir  = make_temp_dir()
    out_dir = make_temp_dir()
    try:
        fname = Path(dwg_path).name
        shutil.copy2(dwg_path, os.path.join(in_dir, fname))
        cmd = [oda, in_dir, out_dir, "ACAD2018", "DXF", "0", "1"]
        run_cmd(cmd, timeout=120)
        time.sleep(1)
        dxf_files = list(Path(out_dir).rglob("*.dxf")) + list(Path(out_dir).rglob("*.DXF"))
        if not dxf_files:
            return None, "ODA produced no DXF"
        content = dxf_files[0].read_text(encoding='utf-8', errors='replace')
        tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.dxf',
                                          encoding='utf-8', delete=False)
        tmp.write(content); tmp.close()
        try:
            doc = _ezdxf.readfile(tmp.name)
        except:
            doc, _ = _recover.read(tmp.name)
        os.unlink(tmp.name)
        return doc, None
    finally:
        safe_rmtree(in_dir); safe_rmtree(out_dir)


def _analyze_dxf(doc) -> dict:
    """Extract structured data from a DXF modelspace for comparison."""
    msp = doc.modelspace()
    info = {
        'total': 0,
        'by_type': {},
        'by_layer': {},
        'layers': [],
        'texts': [],
        'blocks_used': {},
        'circles': [],
        'lines_count': 0,
        'dimensions': [],
        'hatches': 0,
        'block_defs': [],
    }

    for e in msp:
        info['total'] += 1
        t = e.dxftype()
        info['by_type'][t] = info['by_type'].get(t, 0) + 1
        layer = e.dxf.get('layer', '0')
        info['by_layer'][layer] = info['by_layer'].get(layer, 0) + 1

        if t in ('TEXT', 'MTEXT'):
            try:
                txt = e.dxf.text if t == 'TEXT' else e.plain_mtext()
                if txt and txt.strip():
                    info['texts'].append(txt.strip()[:80])
            except: pass
        elif t == 'INSERT':
            name = e.dxf.get('name', '?')
            info['blocks_used'][name] = info['blocks_used'].get(name, 0) + 1
        elif t == 'CIRCLE':
            info['circles'].append(round(e.dxf.radius, 2))
        elif t == 'LINE':
            info['lines_count'] += 1
        elif t == 'DIMENSION':
            try: info['dimensions'].append(round(e.dxf.actual_measurement, 2))
            except: pass
        elif t == 'HATCH':
            info['hatches'] += 1

    info['layers'] = sorted(set(info['by_layer'].keys()))

    # Block definitions
    try:
        info['block_defs'] = [b.name for b in doc.blocks
                               if not b.name.startswith('*')]
    except: pass

    return info


def _compare_cad(path1, path2, ext1, ext2) -> dict:
    """Compare two CAD files (DXF or DWG)."""
    if ext1 == 'dwg':
        doc1, err1 = _load_dxf_from_path_via_oda(path1)
    else:
        doc1, err1 = _load_dxf_for_compare(path1, ext1)

    if ext2 == 'dwg':
        doc2, err2 = _load_dxf_from_path_via_oda(path2)
    else:
        doc2, err2 = _load_dxf_for_compare(path2, ext2)

    if doc1 is None:
        return {"success": False, "error": f"Cannot read file 1: {err1}"}
    if doc2 is None:
        return {"success": False, "error": f"Cannot read file 2: {err2}"}

    a1 = _analyze_dxf(doc1)
    a2 = _analyze_dxf(doc2)

    changes = []
    summary_stats = {
        'file1_entities': a1['total'],
        'file2_entities': a2['total'],
        'file1_layers':   len(a1['layers']),
        'file2_layers':   len(a2['layers']),
    }

    # ── Entity type changes ─────────────────────────────────
    all_types = set(list(a1['by_type'].keys()) + list(a2['by_type'].keys()))

    # Map entity types to human-readable names
    TYPE_NAMES = {
        'LINE': 'Line', 'CIRCLE': 'Circle', 'ARC': 'Arc',
        'LWPOLYLINE': 'Polyline', 'POLYLINE': 'Polyline',
        'TEXT': 'Text Label', 'MTEXT': 'Text Label',
        'DIMENSION': 'Dimension', 'INSERT': 'Block Reference',
        'HATCH': 'Hatch/Fill', 'SPLINE': 'Spline Curve',
        'ELLIPSE': 'Ellipse', 'SOLID': 'Solid Fill',
        'POINT': 'Point', 'LEADER': 'Leader/Arrow',
        'IMAGE': 'Raster Image', 'XREF': 'External Reference',
    }

    for t in sorted(all_types):
        c1 = a1['by_type'].get(t, 0)
        c2 = a2['by_type'].get(t, 0)
        if c1 == c2: continue
        diff = c2 - c1
        name = TYPE_NAMES.get(t, t)
        severity = 'high' if abs(diff) >= 5 else 'medium' if abs(diff) >= 2 else 'low'
        changes.append({
            'category': 'GEOMETRY',
            'severity': severity,
            'type': 'added' if diff > 0 else 'removed',
            'icon': '➕' if diff > 0 else '➖',
            'change': f"{'Added' if diff > 0 else 'Removed'} {abs(diff)} {name}{'s' if abs(diff)>1 else ''}",
            'detail': f"File 1: {c1}  →  File 2: {c2}",
        })

    # ── Block reference changes (DOOR, WINDOW, etc.) ─────────
    all_blocks = set(list(a1['blocks_used'].keys()) + list(a2['blocks_used'].keys()))
    for block in sorted(all_blocks):
        c1 = a1['blocks_used'].get(block, 0)
        c2 = a2['blocks_used'].get(block, 0)
        if c1 == c2: continue
        diff = c2 - c1
        # Guess what the block is from its name
        bname = block.upper()
        kind = ('door' if any(k in bname for k in ['DOOR','DR','DOR']) else
                'window' if any(k in bname for k in ['WIND','WIN','WN','WDW']) else
                'column' if any(k in bname for k in ['COL','COLUMN','PILLAR']) else
                'stair' if any(k in bname for k in ['STAIR','STEP']) else
                'furniture' if any(k in bname for k in ['FURN','CHAIR','TABLE','DESK']) else
                f"'{block}' symbol")
        changes.append({
            'category': 'SYMBOL',
            'severity': 'high',
            'type': 'added' if diff > 0 else 'removed',
            'icon': '🚪' if 'door' in kind else '🪟' if 'window' in kind else '🏛' if 'column' in kind else '🔷',
            'change': f"{'Added' if diff > 0 else 'Removed'} {abs(diff)} {kind}",
            'detail': f"Block '{block}': {c1} → {c2}",
        })

    # ── Layer changes ────────────────────────────────────────
    new_layers = sorted(set(a2['layers']) - set(a1['layers']))
    removed_layers = sorted(set(a1['layers']) - set(a2['layers']))
    for l in new_layers:
        n = a2['by_layer'].get(l, 0)
        changes.append({
            'category': 'LAYER',
            'severity': 'medium',
            'type': 'added',
            'icon': '📋',
            'change': f"New layer added: '{l}'",
            'detail': f"{n} entities on this layer",
        })
    for l in removed_layers:
        changes.append({
            'category': 'LAYER',
            'severity': 'medium',
            'type': 'removed',
            'icon': '📋',
            'change': f"Layer removed: '{l}'",
            'detail': f"Layer with {a1['by_layer'].get(l,0)} entities no longer present",
        })

    # ── Text label changes ───────────────────────────────────
    t1_set = set(a1['texts'])
    t2_set = set(a2['texts'])
    for t in sorted(t2_set - t1_set)[:10]:
        changes.append({
            'category': 'TEXT',
            'severity': 'low',
            'type': 'added',
            'icon': '📝',
            'change': f"New text label: \"{t}\"",
            'detail': 'Text annotation added',
        })
    for t in sorted(t1_set - t2_set)[:10]:
        changes.append({
            'category': 'TEXT',
            'severity': 'low',
            'type': 'removed',
            'icon': '📝',
            'change': f"Removed text label: \"{t}\"",
            'detail': 'Text annotation removed',
        })

    # ── Dimension changes ────────────────────────────────────
    dc1 = a1['by_type'].get('DIMENSION', 0)
    dc2 = a2['by_type'].get('DIMENSION', 0)
    if dc1 != dc2:
        diff = dc2 - dc1
        changes.append({
            'category': 'DIMENSION',
            'severity': 'medium',
            'type': 'added' if diff > 0 else 'removed',
            'icon': '📐',
            'change': f"{'Added' if diff>0 else 'Removed'} {abs(diff)} dimension annotation{'s' if abs(diff)>1 else ''}",
            'detail': f"{dc1} → {dc2} dimensions",
        })

    # ── Overall summary ──────────────────────────────────────
    total_diff = a2['total'] - a1['total']
    overall = ('identical' if not changes else
               'minor differences' if len(changes) <= 2 else
               'moderate changes' if len(changes) <= 6 else
               'significant changes')

    return {
        'success': True,
        'fileType': 'CAD Drawing',
        'overall': overall,
        'totalChanges': len(changes),
        'added': sum(1 for c in changes if c['type'] == 'added'),
        'removed': sum(1 for c in changes if c['type'] == 'removed'),
        'changes': changes,
        'stats': {**summary_stats,
                  'entity_diff': total_diff,
                  'file1_blocks': len(a1['blocks_used']),
                  'file2_blocks': len(a2['blocks_used']),
                  'new_layers': new_layers,
                  'removed_layers': removed_layers},
    }


# ── IFC Comparison ────────────────────────────────────────────
def _compare_ifc(path1, path2) -> dict:
    try:
        import ifcopenshell
    except ImportError:
        return {"success": False, "error": "ifcopenshell not installed"}

    try:
        ifc1 = ifcopenshell.open(path1)
        ifc2 = ifcopenshell.open(path2)
    except Exception as e:
        return {"success": False, "error": f"Cannot open IFC: {e}"}

    changes = []
    stats = {}

    # Count elements by type
    ifc_types = [
        'IfcWall','IfcWallStandardCase','IfcSlab','IfcRoof','IfcColumn',
        'IfcBeam','IfcDoor','IfcWindow','IfcStair','IfcRamp',
        'IfcFurnishingElement','IfcSpace','IfcBuildingStorey',
        'IfcFlowTerminal','IfcFlowSegment','IfcOpeningElement',
        'IfcPlate','IfcMember','IfcCovering',
    ]

    ifc_icons = {
        'IfcWall': '🧱', 'IfcWallStandardCase': '🧱', 'IfcSlab': '⬜',
        'IfcRoof': '🏠', 'IfcColumn': '🏛', 'IfcBeam': '━',
        'IfcDoor': '🚪', 'IfcWindow': '🪟', 'IfcStair': '🪜',
        'IfcSpace': '📐', 'IfcFurnishingElement': '🪑',
        'IfcBuildingStorey': '🏢', 'IfcFlowTerminal': '💡',
    }

    for ifc_type in ifc_types:
        try:
            c1 = len(ifc1.by_type(ifc_type))
            c2 = len(ifc2.by_type(ifc_type))
        except:
            continue
        if c1 == c2: continue
        diff = c2 - c1
        human = ifc_type.replace('Ifc','').replace('StandardCase','')
        changes.append({
            'category': 'IFC_ELEMENT',
            'severity': 'high' if abs(diff) >= 3 else 'medium',
            'type': 'added' if diff > 0 else 'removed',
            'icon': ifc_icons.get(ifc_type, '🔷'),
            'change': f"{'Added' if diff>0 else 'Removed'} {abs(diff)} {human}{'s' if abs(diff)>1 else ''}",
            'detail': f"{c1} → {c2}",
        })

    # Storey changes
    def get_storeys(ifc):
        try:
            return {s.Name: s for s in ifc.by_type('IfcBuildingStorey')}
        except: return {}

    s1 = get_storeys(ifc1); s2 = get_storeys(ifc2)
    for name in set(s2.keys()) - set(s1.keys()):
        changes.append({'category':'STOREY','severity':'high','type':'added',
                        'icon':'🏢','change':f"New building storey: '{name}'",
                        'detail':'New floor/level added to model'})
    for name in set(s1.keys()) - set(s2.keys()):
        changes.append({'category':'STOREY','severity':'high','type':'removed',
                        'icon':'🏢','change':f"Storey removed: '{name}'",
                        'detail':'Floor/level removed from model'})

    # Total element counts
    def total_products(ifc):
        try: return len(ifc.by_type('IfcProduct'))
        except: return 0

    tp1 = total_products(ifc1); tp2 = total_products(ifc2)
    stats = {
        'file1_elements': tp1, 'file2_elements': tp2,
        'file1_storeys': len(s1), 'file2_storeys': len(s2),
        'schema1': ifc1.schema, 'schema2': ifc2.schema,
    }

    overall = ('identical' if not changes else
               'minor differences' if len(changes) <= 2 else
               'moderate changes' if len(changes) <= 6 else
               'significant changes')

    return {
        'success': True, 'fileType': 'IFC BIM Model',
        'overall': overall,
        'totalChanges': len(changes),
        'added': sum(1 for c in changes if c['type'] == 'added'),
        'removed': sum(1 for c in changes if c['type'] == 'removed'),
        'changes': changes, 'stats': stats,
    }


# ── Document/Text Comparison (PDF, Office) ───────────────────
def _extract_text(path: str, ext: str) -> str:
    """
    Extract plain text from PDF or Office file.
    PDF pipeline: pdfplumber → pypdf → LibreOffice → OCR (tesseract)
    Office pipeline: LibreOffice text export → python-docx fallback
    """
    if ext == 'pdf':

        # ── Method 1: pdfplumber (best for structured PDFs) ──────
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    # Try text extraction
                    t = page.extract_text(x_tolerance=2, y_tolerance=2)
                    if t and t.strip():
                        pages.append(t.strip())
                    # Also extract tables as text
                    try:
                        for table in page.extract_tables():
                            for row in table:
                                row_text = '  |  '.join(str(c or '') for c in row)
                                if row_text.strip():
                                    pages.append(row_text)
                    except Exception:
                        pass
            text = '\n'.join(pages)
            if text.strip():
                print(f"[PDF] pdfplumber: {len(text)} chars from {len(pages)} pages", flush=True)
                return text
            print("[PDF] pdfplumber returned empty — trying pypdf", flush=True)
        except Exception as e:
            print(f"[PDF] pdfplumber error: {e}", flush=True)

        # ── Method 2: pypdf ──────────────────────────────────────
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            # Check if PDF is encrypted
            if reader.is_encrypted:
                try:
                    reader.decrypt('')  # try empty password
                except Exception:
                    print("[PDF] PDF is password-protected", flush=True)
                    return ""
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t.strip())
            text = '\n'.join(pages)
            if text.strip():
                print(f"[PDF] pypdf: {len(text)} chars", flush=True)
                return text
            print("[PDF] pypdf returned empty — likely scanned PDF", flush=True)
        except Exception as e:
            print(f"[PDF] pypdf error: {e}", flush=True)

        # ── Method 3: pdftotext CLI (poppler) ────────────────────
        try:
            rc, stdout, stderr = run_cmd(['pdftotext', '-layout', path, '-'], timeout=30)
            if rc == 0 and stdout.strip():
                print(f"[PDF] pdftotext: {len(stdout)} chars", flush=True)
                return stdout
        except Exception:
            pass

        # ── Method 4: OCR via pypdfium2 + Tesseract (Apache 2.0) ──
        # pypdfium2 replaces PyMuPDF (AGPL) — no Poppler needed, commercial-safe
        try:
            import pypdfium2 as pdfium
            import pytesseract
            from PIL import Image

            tess = find_tesseract()
            if not tess:
                raise FileNotFoundError("Tesseract not found")
            pytesseract.pytesseract.tesseract_cmd = tess

            print("[PDF] OCR with pypdfium2 + Tesseract...", flush=True)
            pdf_doc = pdfium.PdfDocument(path)
            pages = []
            max_pages = min(len(pdf_doc), 10)
            for i in range(max_pages):
                pg = pdf_doc[i]
                bitmap = pg.render(scale=2)   # 2x ≈ 144dpi
                img = bitmap.to_pil()
                t = pytesseract.image_to_string(img, lang='eng')
                if t and t.strip():
                    pages.append(t.strip())
            pdf_doc.close()
            text = '\n'.join(pages)
            if text.strip():
                print(f"[PDF] OCR: {len(text)} chars from {max_pages} pages", flush=True)
                return text
            print("[PDF] OCR returned empty — PDF may be blank", flush=True)
        except ImportError as e:
            print(f"[PDF] OCR import error: {e} — run: pip install pypdfium2 pytesseract Pillow", flush=True)
        except FileNotFoundError as e:
            print(f"[PDF] Tesseract not found: {e}", flush=True)
        except Exception as e:
            print(f"[PDF] OCR error: {e}", flush=True)

        # ── Method 5: LibreOffice text export ────────────────────
        lo = find_libreoffice()
        if lo:
            out_dir = make_temp_dir()
            try:
                run_cmd([lo, '--headless', '--norestore', '--convert-to', 'txt:Text',
                         '--outdir', out_dir, path], timeout=60)
                txt_files = list(Path(out_dir).glob('*.txt'))
                if txt_files:
                    text = txt_files[0].read_text(encoding='utf-8', errors='replace')
                    if text.strip():
                        print(f"[PDF] LibreOffice: {len(text)} chars", flush=True)
                        return text
            finally:
                safe_rmtree(out_dir)

        print("[PDF] All extraction methods failed", flush=True)
        return ""

    elif ext in OFFICE_EXTS:
        # ── LibreOffice text export ───────────────────────────────
        lo = find_libreoffice()
        if lo:
            out_dir = make_temp_dir()
            try:
                run_cmd([lo, '--headless', '--norestore', '--convert-to', 'txt:Text',
                         '--outdir', out_dir, path], timeout=60)
                txt_files = list(Path(out_dir).glob('*.txt'))
                if txt_files:
                    text = txt_files[0].read_text(encoding='utf-8', errors='replace')
                    if text.strip():
                        return text
            finally:
                safe_rmtree(out_dir)

        # ── python-docx fallback for .docx ────────────────────────
        if ext == 'docx':
            try:
                import docx
                doc = docx.Document(path)
                return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                pass

        # ── pdfplumber fallback for PDF-like office formats ───────
        if ext in ('xlsx', 'xls'):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
                rows = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        r = '  |  '.join(str(c) for c in row if c is not None)
                        if r.strip():
                            rows.append(r)
                return '\n'.join(rows)
            except ImportError:
                pass

    return ""


def _compare_documents(path1, path2, ext1, ext2) -> dict:
    import difflib

    print(f"[DOC] Extracting text from: {Path(path1).name}", flush=True)
    text1 = _extract_text(path1, ext1)
    print(f"[DOC] Extracting text from: {Path(path2).name}", flush=True)
    text2 = _extract_text(path2, ext2)
    print(f"[DOC] Text lengths: {len(text1)}, {len(text2)}", flush=True)

    # ── Get PDF metadata (page count, size) for fallback comparison ──
    def pdf_meta(path):
        meta = {'pages': 0, 'size_kb': round(os.path.getsize(path) / 1024, 1)}
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            meta['pages'] = len(reader.pages)
            meta['encrypted'] = reader.is_encrypted
        except Exception:
            pass
        return meta

    m1 = pdf_meta(path1) if ext1 == 'pdf' else {}
    m2 = pdf_meta(path2) if ext2 == 'pdf' else {}

    # ── Both failed — return metadata diff + helpful guidance ────
    if not text1 and not text2:
        changes = []
        if m1 and m2:
            if m1['pages'] != m2['pages']:
                diff = m2['pages'] - m1['pages']
                changes.append({
                    'category': 'STRUCTURE', 'severity': 'high',
                    'type': 'added' if diff > 0 else 'removed',
                    'icon': '📄',
                    'change': f"Page count changed: {m1['pages']} → {m2['pages']} pages",
                    'detail': f"{'Added' if diff>0 else 'Removed'} {abs(diff)} page{'s' if abs(diff)>1 else ''}"
                })
            if abs(m1['size_kb'] - m2['size_kb']) > 10:
                changes.append({
                    'category': 'STRUCTURE', 'severity': 'medium',
                    'type': 'modified', 'icon': '💾',
                    'change': f"File size changed: {m1['size_kb']} KB → {m2['size_kb']} KB",
                    'detail': f"{'Larger' if m2['size_kb'] > m1['size_kb'] else 'Smaller'} file may indicate content changes"
                })
        return {
            'success': True,
            'fileType': 'PDF (Scanned/Image)',
            'overall': 'metadata only — text not extractable',
            'totalChanges': len(changes),
            'added': 0, 'removed': 0,
            'changes': changes,
            'warning': (
                'These PDFs appear to be scanned images with no text layer. '
                'Only structural metadata (page count, file size) could be compared.\n\n'
                'To enable full text comparison, either:\n'
                '• Install Tesseract OCR: sudo apt install tesseract-ocr (Ubuntu) or https://github.com/UB-Mannheim/tesseract/wiki (Windows)\n'
                '• Or use text-based PDFs instead of scanned images'
            ),
            'stats': {
                'file1_pages': m1.get('pages', '?'), 'file2_pages': m2.get('pages', '?'),
                'file1_size_kb': m1.get('size_kb', 0), 'file2_size_kb': m2.get('size_kb', 0),
            }
        }

    # ── One file failed — still do partial diff ───────────────────
    if not text1 or not text2:
        failed_file = "File 1" if not text1 else "File 2"
        ok_text     = text2 if not text1 else text1
        ok_lines    = len([l for l in ok_text.splitlines() if l.strip()])
        meta_failed = m1 if not text1 else m2

        changes = [{
            'category': 'EXTRACTION', 'severity': 'high',
            'type': 'modified', 'icon': '⚠️',
            'change': f"{failed_file} could not be read as text",
            'detail': (f"Likely a scanned/image PDF ({meta_failed.get('pages','?')} pages, "
                       f"{meta_failed.get('size_kb','?')} KB). "
                       f"Install Tesseract OCR for full comparison.")
        }]

        # Page count diff if available
        if m1.get('pages') and m2.get('pages') and m1['pages'] != m2['pages']:
            diff = m2['pages'] - m1['pages']
            changes.append({
                'category': 'STRUCTURE', 'severity': 'medium',
                'type': 'added' if diff > 0 else 'removed', 'icon': '📄',
                'change': f"Page count: {m1['pages']} → {m2['pages']}",
                'detail': f"{'Added' if diff>0 else 'Removed'} {abs(diff)} page(s)"
            })

        return {
            'success': True,
            'fileType': 'Document (partial)',
            'overall': f'{failed_file} unreadable — partial comparison only',
            'totalChanges': len(changes),
            'added': 0, 'removed': 0,
            'changes': changes,
            'warning': (
                f'{failed_file} appears to be a scanned/image PDF with no text layer.\n'
                'Install Tesseract OCR for full text comparison:\n'
                '  Ubuntu: sudo apt install tesseract-ocr\n'
                '  Windows: https://github.com/UB-Mannheim/tesseract/wiki'
            ),
            'stats': {
                'file1_lines': ok_lines if not text2 else 0,
                'file2_lines': ok_lines if not text1 else 0,
            }
        }

    # ── Both have text — do full diff ──────────────────────────────
    lines1 = [l.strip() for l in text1.splitlines() if l.strip()]
    lines2 = [l.strip() for l in text2.splitlines() if l.strip()]

    changes = []
    diff = list(difflib.ndiff(lines1, lines2))
    added   = [l[2:] for l in diff if l.startswith('+ ')]
    removed = [l[2:] for l in diff if l.startswith('- ')]

    for line in added[:20]:
        changes.append({'category':'CONTENT','severity':'medium','type':'added',
                        'icon':'➕','change':f"Added: \"{line[:100]}\"",
                        'detail':'New content in revised document'})
    for line in removed[:20]:
        changes.append({'category':'CONTENT','severity':'medium','type':'removed',
                        'icon':'➖','change':f"Removed: \"{line[:100]}\"",
                        'detail':'Content removed from original'})

    overall = ('identical' if not changes else
               'minor differences' if len(changes) <= 3 else
               'moderate changes' if len(changes) <= 10 else
               'significant changes')

    return {
        'success': True, 'fileType': 'Document',
        'overall': overall,
        'totalChanges': len(changes),
        'added': len(added), 'removed': len(removed),
        'changes': changes,
        'stats': {'file1_lines': len(lines1), 'file2_lines': len(lines2),
                  'lines_added': len(added), 'lines_removed': len(removed)},
    }


# ── Image Comparison ─────────────────────────────────────────
def _compare_images(path1, path2) -> dict:
    try:
        import numpy as np
        from PIL import Image
    except ImportError:
        return {"success": False, "error": "Pillow not installed. Run: pip install Pillow"}

    try:
        img1 = np.array(Image.open(path1).convert('RGB'), dtype=np.float32)
        img2 = np.array(Image.open(path2).convert('RGB'), dtype=np.float32)
    except Exception as e:
        return {"success": False, "error": f"Cannot open images: {e}"}

    changes = []
    h1, w1 = img1.shape[:2]
    h2, w2 = img2.shape[:2]

    if (h1, w1) != (h2, w2):
        changes.append({'category':'SIZE','severity':'high','type':'modified',
                        'icon':'📐',
                        'change':f"Image dimensions changed: {w1}×{h1} → {w2}×{h2}",
                        'detail':'Different canvas size'})
        # Resize for pixel comparison
        from PIL import Image as PILImage
        img2_resized = np.array(PILImage.open(path2).convert('RGB').resize((w1,h1)), dtype=np.float32)
        diff = np.abs(img1 - img2_resized)
    else:
        diff = np.abs(img1 - img2)

    total_pixels = img1.shape[0] * img1.shape[1]
    changed_pixels = int(np.sum(np.any(diff > 10, axis=2)))
    pct = changed_pixels / total_pixels * 100

    if pct > 0.1:
        severity = 'high' if pct > 20 else 'medium' if pct > 5 else 'low'
        changes.append({'category':'PIXELS','severity':severity,'type':'modified',
                        'icon':'🖼',
                        'change':f"{pct:.1f}% of pixels differ between images",
                        'detail':f"{changed_pixels:,} of {total_pixels:,} pixels changed"})

    overall = ('identical' if pct < 0.1 else
               f'{pct:.1f}% pixel difference')

    return {
        'success': True, 'fileType': 'Image',
        'overall': overall,
        'totalChanges': len(changes),
        'added': 0, 'removed': 0,
        'changes': changes,
        'stats': {'width1':w1,'height1':h1,'width2':w2,'height2':h2,
                  'pixel_diff_pct': round(pct,2)},
    }


# ── Metadata comparison (mixed types) ────────────────────────
def _compare_metadata(path1, path2) -> dict:
    s1 = os.path.getsize(path1)
    s2 = os.path.getsize(path2)
    changes = []
    if s1 != s2:
        diff_pct = abs(s2-s1)/s1*100 if s1 > 0 else 100
        changes.append({'category':'SIZE','severity':'medium','type':'modified',
                        'icon':'💾',
                        'change':f"File size changed by {diff_pct:.1f}%",
                        'detail':f"{s1:,} bytes → {s2:,} bytes"})
    return {
        'success': True, 'fileType': 'File',
        'overall': 'different file sizes' if changes else 'same size',
        'totalChanges': len(changes),
        'added': 0, 'removed': 0,
        'changes': changes,
        'stats': {'size1': s1, 'size2': s2},
    }
# ══════════════════════════════════════════════════════════════════
#  PHASE 3: PDF REDACTION
# ══════════════════════════════════════════════════════════════════

# Render DPI for pages that get rasterized. Must be passed to BOTH the
# render call and the save call — see _rewrite_pdf_pages.
REDACTION_DPI = 150
FLATTEN_DPI   = {"screen": 150, "print": 300}


def redact_pdf(path: str, regions: list, output: str = "", burn: bool = True) -> dict:
    """
    Redact rectangular regions from a PDF.
    burn=True  → permanently destroy content (production)
    burn=False → add black rectangle overlay only (preview)

    regions: [{page, x, y, width, height, reason}]
    Coordinates are in PDF points (72 pt = 1 inch).
    """
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}
    if not regions:
        return {"success": False, "error": "No redaction regions provided"}

    try:
        import pypdfium2 as pdfium

        # Drop regions pointing outside the document rather than failing the
        # whole request over one bad index.
        doc = pdfium.PdfDocument(path)
        page_count = len(doc)
        doc.close()

        by_page = {
            idx: items
            for idx, items in _group_by_page(regions).items()
            if 0 <= idx < page_count
        }
        if not by_page:
            return {"success": False,
                    "error": f"No regions fall within this document's {page_count} page(s)"}

        if not output:
            base = os.path.splitext(path)[0]
            output = base + ("_redacted.pdf" if burn else "_redact_preview.pdf")

        redacted_pages = 0
        if burn:
            redacted_pages = _rewrite_pdf_pages(
                path, output, REDACTION_DPI, by_page, _draw_redactions
            )
        # Preview mode draws the overlay client-side; nothing to write here.

        return {
            "success": True,
            "outputPath": output,
            "redactedPages": redacted_pages,
            "totalPages": page_count,
            "totalRegions": len(regions),
            "mode": "permanent" if burn else "preview"
        }

    except ImportError as e:
        return {"success": False, "error": f"Required library missing: {e}. Run: pip install pypdfium2 pypdf Pillow"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _group_by_page(items: list, key: str = "page") -> dict:
    """Group edit items by 0-based page index."""
    grouped = {}
    for item in items:
        idx = int(item.get(key, 1)) - 1
        grouped.setdefault(idx, []).append(item)
    return grouped


def _rewrite_pdf_pages(path: str, output: str, dpi: int,
                       edits_by_page: dict, draw_page) -> int:
    """
    Rebuild a PDF, rasterizing ONLY the pages that carry edits and copying
    every other page through untouched.

    Copying untouched pages matters for correctness, not just speed:
    rasterizing a page discards its vector content and its text layer, so
    blanket-rasterizing silently makes an ENTIRE document unsearchable
    because one region on one page was redacted.

    The render DPI and the DPI declared when saving must be the same number,
    or every rasterized page is rescaled by their ratio — rendering at
    scale=2 (144dpi) while declaring resolution=150 shrank every page to
    96% of its original size.

    draw_page(draw, items, scale_x, scale_y, page_width, page_height)
    performs the actual drawing. Returns the count of rasterized pages.
    """
    import pypdfium2 as pdfium
    from PIL import ImageDraw
    import pypdf
    import io

    src     = pdfium.PdfDocument(path)
    reader  = pypdf.PdfReader(path)
    writer  = pypdf.PdfWriter()
    touched = 0

    try:
        for i in range(len(src)):
            items = edits_by_page.get(i)
            if not items:
                writer.add_page(reader.pages[i])
                continue

            page   = src[i]
            pw, ph = page.get_width(), page.get_height()
            img    = page.render(scale=dpi / 72.0).to_pil().convert("RGB")
            iw, ih = img.size

            draw_page(ImageDraw.Draw(img, "RGBA"), items, iw / pw, ih / ph, pw, ph)

            buf = io.BytesIO()
            img.save(buf, "PDF", resolution=dpi)
            buf.seek(0)
            writer.add_page(pypdf.PdfReader(buf).pages[0])
            touched += 1
    finally:
        src.close()

    with open(output, "wb") as fh:
        writer.write(fh)
    return touched


def _draw_redactions(draw, regions, scale_x, scale_y, page_w, page_h):
    """Burn opaque black boxes over each region (PDF points, origin bottom-left)."""
    for region in regions:
        rx = float(region.get("x", 0))
        ry = float(region.get("y", 0))
        rw = float(region.get("width", 100))
        rh = float(region.get("height", 20))
        draw.rectangle(
            [int(rx * scale_x),        int((page_h - ry - rh) * scale_y),
             int((rx + rw) * scale_x), int((page_h - ry) * scale_y)],
            fill=(0, 0, 0)
        )


def _draw_shapes(draw, shapes, scale_x, scale_y, page_w, page_h):
    """
    Paint annotation shapes onto a rasterized page. Shape coordinates are
    top-left origin (the frontend's screen space at zoom 1, which equals PDF
    points), unlike redaction regions.
    """
    from PIL.ImageColor import getrgb

    for shape in shapes:
        tool  = shape.get("tool", "rect")
        try:
            rgb = getrgb(shape.get("color", "#FF0000"))
        except Exception:
            rgb = (255, 0, 0)
        width = max(1, int(float(shape.get("strokeWidth", 2)) * scale_x * 0.5))

        if tool in ("line", "arrow", "dimension"):
            draw.line([float(shape.get("x1", 0)) * scale_x, float(shape.get("y1", 0)) * scale_y,
                       float(shape.get("x2", 0)) * scale_x, float(shape.get("y2", 0)) * scale_y],
                      fill=rgb, width=width)

        elif tool in ("rect", "highlight", "redact", "ellipse",
                      "underline", "strikeout", "squiggly"):
            x = float(shape.get("x", 0)) * scale_x
            y = float(shape.get("y", 0)) * scale_y
            w = float(shape.get("width", 100)) * scale_x
            h = float(shape.get("height", 50)) * scale_y
            if tool == "redact":
                draw.rectangle([x, y, x + w, y + h], fill=(0, 0, 0))
            elif tool == "ellipse":
                draw.ellipse([x, y, x + w, y + h], outline=rgb, width=width)
            elif tool == "underline":
                draw.line([x, y + h, x + w, y + h], fill=rgb, width=width)
            elif tool == "strikeout":
                draw.line([x, y + h / 2, x + w, y + h / 2], fill=rgb, width=width)
            elif tool == "squiggly":
                draw.line([x, y + h, x + w, y + h], fill=rgb, width=width)
            else:
                opacity = int(float(shape.get("opacity", 0.15)) * 255)
                draw.rectangle([x, y, x + w, y + h],
                               outline=rgb, width=width, fill=(*rgb, opacity))

        elif tool == "circle":
            cx = float(shape.get("cx", 0)) * scale_x
            cy = float(shape.get("cy", 0)) * scale_y
            r  = float(shape.get("r", 50)) * scale_x
            draw.ellipse([cx - r, cy - r, cx + r, cy + r], outline=rgb, width=width)

        elif tool in ("freehand", "cloud", "polygon", "polyline"):
            pts = shape.get("points", [])
            if len(pts) >= 2:
                scaled = [(p["x"] * scale_x, p["y"] * scale_y) for p in pts]
                if tool == "polygon":
                    scaled.append(scaled[0])   # close the ring
                draw.line(scaled, fill=rgb, width=width)

        elif tool in ("text", "stamp", "callout", "note"):
            text = shape.get("text", "")
            if text:
                draw.text((float(shape.get("x", 0)) * scale_x,
                           float(shape.get("y", 0)) * scale_y), text, fill=rgb)


# ══════════════════════════════════════════════════════════════════
#  OCR: SCANNED PDF -> SEARCHABLE PDF
# ══════════════════════════════════════════════════════════════════

def _page_has_text(page, min_chars: int = 10) -> bool:
    """
    True if a pypdfium2 page already carries an extractable text layer.
    Used to skip pages that don't need OCR — re-OCRing a digital page
    would rasterize it and *lose* quality and real text for no gain.
    """
    try:
        textpage = page.get_textpage()
        text     = textpage.get_text_range() or ""
        textpage.close()
        return len(text.strip()) >= min_chars
    except Exception:
        return False


def ocr_pdf_to_searchable(path: str, output: str = "", lang: str = "eng",
                          dpi: int = 300, skip_text_pages: bool = True) -> dict:
    """
    Produce a searchable PDF from a scanned/image PDF by adding an invisible
    text layer, leaving the visible page appearance unchanged.

    Uses Tesseract's own PDF renderer (image_to_pdf_or_hocr) rather than
    positioning invisible text by hand — Tesseract already knows each glyph's
    exact box from recognition, so its output aligns text to the image
    correctly, which hand-placed word boxes reliably get subtly wrong.

    Pages that already have a text layer are copied through untouched
    (skip_text_pages=True), so this is safe to run on a mixed document.
    """
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}

    tess = find_tesseract()
    if not tess:
        return {"success": False, "error": "TESSERACT_NOT_FOUND",
                "hint": "Install Tesseract OCR: 'sudo apt install tesseract-ocr' "
                        "(Linux) or https://github.com/UB-Mannheim/tesseract/wiki (Windows)"}

    # Guard rails: dpi below ~150 wrecks recognition accuracy, above 600 costs
    # a lot of time/memory for no measurable gain.
    dpi = max(150, min(int(dpi or 300), 600))

    try:
        import pypdfium2 as pdfium
        import pytesseract
        import pypdf
        import io

        pytesseract.pytesseract.tesseract_cmd = tess

        if not output:
            output = os.path.splitext(path)[0] + "_searchable.pdf"

        src        = pdfium.PdfDocument(path)
        reader     = pypdf.PdfReader(path)
        writer     = pypdf.PdfWriter()
        total      = len(src)
        ocr_pages  = 0
        kept_pages = 0

        print(f"[OCR] {total} page(s) @ {dpi}dpi lang={lang}", flush=True)

        for i in range(total):
            page = src[i]

            if skip_text_pages and _page_has_text(page):
                writer.add_page(reader.pages[i])
                kept_pages += 1
                continue

            # pypdfium2's render scale is relative to 72dpi
            bitmap = page.render(scale=dpi / 72.0)
            img    = bitmap.to_pil().convert("RGB")

            # --dpi must be passed explicitly: Tesseract otherwise assumes
            # 70dpi for a bitmap carrying no resolution metadata and sizes
            # the output page from that, which silently inflated every
            # OCR'd page to (dpi/70)x its true dimensions.
            page_pdf = pytesseract.image_to_pdf_or_hocr(
                img, lang=lang, extension="pdf", config=f"--dpi {dpi}"
            )
            ocr_page = pypdf.PdfReader(io.BytesIO(page_pdf))
            for p in ocr_page.pages:
                writer.add_page(p)
            ocr_pages += 1

        src.close()

        with open(output, "wb") as fh:
            writer.write(fh)

        print(f"[OCR] done — {ocr_pages} OCR'd, {kept_pages} already searchable", flush=True)

        return {
            "success": True,
            "outputPath": output,
            "totalPages": total,
            "ocrPages": ocr_pages,
            "skippedPages": kept_pages,
            "language": lang,
            "dpi": dpi,
        }

    except ImportError as e:
        return {"success": False,
                "error": f"Required library missing: {e}",
                "hint": "pip install pypdfium2 pytesseract pypdf Pillow"}
    except pytesseract.TesseractError as e:
        # Most common cause: requested language pack not installed
        return {"success": False, "error": f"Tesseract failed: {e}",
                "hint": f"Is the '{lang}' language pack installed? "
                        f"e.g. 'sudo apt install tesseract-ocr-{lang}'"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ══════════════════════════════════════════════════════════════════
#  PHASE 3: FLATTEN ANNOTATIONS TO PDF
# ══════════════════════════════════════════════════════════════════

def flatten_annotations_to_pdf(path: str, shapes: list, output: str = "", quality: str = "screen") -> dict:
    """
    Flatten annotation shapes onto a PDF by rendering pages to images
    and drawing the shapes using PIL.

    shapes: [{tool, pageNumber, color, strokeWidth, x1,y1,x2,y2, ...}]
    """
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}

    try:
        if not output:
            base   = os.path.splitext(path)[0]
            output = base + "_annotated.pdf"

        dpi     = FLATTEN_DPI.get(quality, FLATTEN_DPI["screen"])
        by_page = _group_by_page(shapes, key="pageNumber")

        flattened_pages = _rewrite_pdf_pages(
            path, output, dpi, by_page, _draw_shapes
        )

        return {
            "success":        True,
            "outputPath":     output,
            "flattenedPages": flattened_pages,
            "shapes":         len(shapes),
            "dpi":            dpi
        }

    except ImportError as e:
        return {"success": False, "error": f"Missing library: {e}. Run: pip install pypdfium2 pypdf Pillow"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ══════════════════════════════════════════════════════════════════
#  PHASE 3: PDF FORM FILLING
# ══════════════════════════════════════════════════════════════════

# AcroForm field flags (PDF 32000-1, tables 226/228/230). Stored as a single
# /Ff integer whose bits mean different things per field type.
FF_READ_ONLY   = 1 << 0
FF_REQUIRED    = 1 << 1
FF_MULTILINE   = 1 << 12   # text
FF_PASSWORD    = 1 << 13   # text
FF_RADIO       = 1 << 15   # button
FF_PUSHBUTTON  = 1 << 16   # button
FF_COMBO       = 1 << 17   # choice
FF_MULTISELECT = 1 << 21   # choice

TRUTHY = {"true", "yes", "on", "1", "y", "checked"}


def _form_field_kind(field_type: str, flags: int) -> str:
    """Map a raw /FT plus its flag bits onto a control the UI can render."""
    if field_type == "/Btn":
        if flags & FF_PUSHBUTTON: return "button"
        if flags & FF_RADIO:      return "radio"
        return "checkbox"
    if field_type == "/Ch":
        return "dropdown" if flags & FF_COMBO else "listbox"
    if field_type == "/Sig":
        return "signature"
    if flags & FF_PASSWORD:  return "password"
    if flags & FF_MULTILINE: return "textarea"
    return "text"


def _choice_options(field) -> list:
    """
    Normalise /Opt into {value,label} pairs. Entries are either a bare
    string, or an [export_value, display_label] pair.
    """
    options = []
    for opt in (field.get("/Opt") or []):
        if isinstance(opt, (list, tuple)):
            export  = str(opt[0])
            display = str(opt[1]) if len(opt) > 1 else export
        else:
            export = display = str(opt)
        options.append({"value": export, "label": display})
    return options


def _checkbox_states(field) -> tuple:
    """
    A checkbox's "on" value is whatever its appearance dictionary calls it —
    commonly /Yes but just as legitimately /On or /1 — so it has to be read
    from the field rather than assumed.
    """
    states = [str(s) for s in (field.get("/_States_") or [])
              if not isinstance(s, (list, tuple))]
    on_state = next((s for s in states if s != "/Off"), "/Yes")
    return on_state, "/Off"


def _qualified_field_name(node) -> str:
    """
    Build a field's fully qualified name (parent.child) by walking up the
    /Parent chain, matching the keys PdfReader.get_fields() returns.
    """
    parts, guard = [], 0
    while node is not None and guard < 32:
        title = node.get("/T")
        if title is not None:
            parts.append(str(title))
        parent = node.get("/Parent")
        node = parent.get_object() if parent is not None else None
        guard += 1
    return ".".join(reversed(parts))


def _field_page_numbers(reader) -> dict:
    """Map field name -> 1-based page number of the widget that renders it."""
    pages = {}
    for index, page in enumerate(reader.pages, start=1):
        for annot in (page.get("/Annots") or []):
            try:
                name = _qualified_field_name(annot.get_object())
            except Exception:
                continue
            if name and name not in pages:
                pages[name] = index
    return pages


def _is_truthy(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lstrip("/").lower() in TRUTHY


def _strip_form_interactivity(writer) -> None:
    """
    Remove widget annotations and the AcroForm dictionary.

    pypdf's flatten=True paints each field's value into the page content
    stream but leaves the interactive field in place, so the result still
    opens as an editable form. Dropping the widgets and the AcroForm makes
    the painted values the only remaining representation.
    """
    from pypdf.generic import NameObject, ArrayObject

    for page in writer.pages:
        annots = page.get("/Annots")
        if not annots:
            continue
        kept = ArrayObject([
            a for a in annots
            if str(a.get_object().get("/Subtype", "")) != "/Widget"
        ])
        if len(kept):
            page[NameObject("/Annots")] = kept
        elif "/Annots" in page:
            del page[NameObject("/Annots")]

    root = writer.root_object
    if "/AcroForm" in root:
        del root[NameObject("/AcroForm")]


def fill_pdf_form(path: str, fields: dict, output: str = "", flatten: bool = False) -> dict:
    """
    Fill PDF AcroForm fields.
    fields: {"FieldName": "Value", ...}

    Values are coerced per field type — checkboxes and radios need their
    appearance-state name (/Yes, /On, ...), not a stringified boolean.
    """
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}

    try:
        from pypdf import PdfReader, PdfWriter

        reader      = PdfReader(path)
        form_fields = reader.get_fields() or {}
        if not form_fields:
            return {"success": False,
                    "error": "This PDF has no fillable form fields (no AcroForm found)"}

        writer = PdfWriter(clone_from=path)

        resolved, skipped = {}, {}
        for name, raw_value in (fields or {}).items():
            field = form_fields.get(name)
            if field is None:
                skipped[name] = "no such field in this PDF"
                continue

            flags = int(field.get("/Ff", 0) or 0)
            if flags & FF_READ_ONLY:
                skipped[name] = "field is read-only"
                continue

            kind = _form_field_kind(str(field.get("/FT", "/Tx")), flags)
            if kind in ("checkbox", "radio"):
                on_state, off_state = _checkbox_states(field)
                resolved[name] = on_state if _is_truthy(raw_value) else off_state
            else:
                resolved[name] = "" if raw_value is None else str(raw_value)

        if resolved:
            # Apply across ALL pages. Updating only page 0 left every field on
            # a later page blank while still reporting it as filled.
            writer.update_page_form_field_values(
                list(writer.pages), resolved, auto_regenerate=False, flatten=flatten
            )
            if flatten:
                _strip_form_interactivity(writer)
            else:
                # Without this many viewers render a filled field as empty,
                # because no appearance stream was generated for the new value.
                writer.set_need_appearances_writer(True)

        if not output:
            base   = os.path.splitext(path)[0]
            output = base + "_filled.pdf"

        with open(output, "wb") as f_out:
            writer.write(f_out)

        return {
            "success":         True,
            "outputPath":      output,
            "filledFields":    resolved,
            "skippedFields":   skipped,
            "availableFields": list(form_fields.keys())
        }

    except ImportError:
        return {"success": False, "error": "pypdf required. Run: pip install pypdf"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# Inspect form fields without filling
def inspect_pdf_form(path: str) -> dict:
    """
    Describe every AcroForm field in enough detail for a client to render
    the right control and validate input: the resolved kind, choice options,
    current value, read-only/required state, max length and page number.
    """
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}

    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        fields = reader.get_fields() or {}
        pages  = _field_page_numbers(reader)

        described = []
        for name, field in fields.items():
            field_type = str(field.get("/FT", "/Tx"))
            flags      = int(field.get("/Ff", 0) or 0)
            kind       = _form_field_kind(field_type, flags)
            raw_value  = field.get("/V", "")

            entry = {
                "name":      name,
                "kind":      kind,
                "type":      field_type,
                "flags":     flags,
                "readOnly":  bool(flags & FF_READ_ONLY),
                "required":  bool(flags & FF_REQUIRED),
                "page":      pages.get(name, 1),
            }

            if kind in ("checkbox", "radio"):
                on_state, _ = _checkbox_states(field)
                entry["onState"] = on_state
                entry["checked"] = str(raw_value) == on_state
                entry["value"]   = str(raw_value or "/Off")
            else:
                entry["value"] = "" if raw_value is None else str(raw_value)

            if kind in ("dropdown", "listbox"):
                entry["options"]     = _choice_options(field)
                entry["multiSelect"] = bool(flags & FF_MULTISELECT)

            if kind in ("text", "textarea", "password"):
                max_len = field.get("/MaxLen")
                if max_len is not None:
                    entry["maxLength"] = int(max_len)
                entry["multiline"] = bool(flags & FF_MULTILINE)

            described.append(entry)

        # Stable, page-then-name ordering so the form UI doesn't reshuffle
        # between requests (dict order follows the PDF's internal layout).
        described.sort(key=lambda f: (f["page"], f["name"]))

        return {
            "success":   True,
            "fields":    described,
            "count":     len(described),
            "pageCount": len(reader.pages)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ══════════════════════════════════════════════════════════════════
#  PHASE 3: IFC MODEL TREE
# ══════════════════════════════════════════════════════════════════

def extract_ifc_tree(path: str) -> dict:
    """
    Extract a hierarchical spatial structure from an IFC file.
    Returns: [{id, name, type, children, expanded, selected, visible}]
    """
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}

    try:
        import ifcopenshell

        ifc = ifcopenshell.open(path)

        def get_children(element) -> list:
            children = []
            try:
                for rel in getattr(element, 'IsDecomposedBy', []):
                    for child in rel.RelatedObjects:
                        children.append(element_to_node(child))
            except Exception:
                pass
            # Also include contained elements for spaces/storeys
            try:
                for rel in getattr(element, 'ContainsElements', []):
                    for child in rel.RelatedElements:
                        children.append(element_to_node(child))
            except Exception:
                pass
            return children

        def element_to_node(element) -> dict:
            name = getattr(element, 'Name', None) or element.is_a()
            return {
                "id":       str(element.GlobalId),
                "name":     str(name),
                "type":     element.is_a(),
                "expanded": False,
                "selected": False,
                "visible":  True,
                "children": get_children(element)
            }

        # Build tree from IfcProject root
        projects = ifc.by_type("IfcProject")
        if not projects:
            return {"success": False, "error": "No IfcProject found in file"}

        tree = [element_to_node(p) for p in projects]
        return {"success": True, "tree": tree, "schema": ifc.schema}

    except ImportError:
        return {"success": False, "error": "ifcopenshell required"}
    except Exception as e:
        return {"success": False, "error": str(e)}



if __name__ == "__main__":
    if not EZDXF_OK:
        print(f'ERROR: ezdxf not installed.\nRun: python -m pip install "ezdxf[draw]"',
              file=sys.stderr)
        sys.exit(1)

    lo    = find_libreoffice()
    oda   = find_oda()
    libre = find_dwg2dxf()

    tess  = find_tesseract()

    # Create C:\Temp on Windows for short paths
    if IS_WINDOWS:
        try:
            os.makedirs("C:\\Temp", exist_ok=True)
            print(f"[INFO] Using C:\\Temp for temp files (avoids path spaces)", flush=True)
        except Exception:
            print("[WARN] Could not create C:\\Temp — using system temp", flush=True)

    print(f"CDE Converter  |  {platform.system()}  |  ezdxf {ezdxf.__version__}  |  port {PORT}")
    print(f"  LibreOffice : {lo   or 'NOT FOUND'}")
    print(f"  ODA         : {oda  or 'NOT FOUND'}")
    print(f"  LibreDWG    : {libre or 'not found'}")
    print(f"  Tesseract   : {tess or 'NOT FOUND — scanned PDF OCR disabled'}")
    print(f"  POST /convert  body: {{\"path\": \"<absolute_path>\", \"contentType\": \"<mime>\"}}")
    print(flush=True)

    HTTPServer(("0.0.0.0", PORT), Handler).serve_forever()


